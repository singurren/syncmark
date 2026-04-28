from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path('/mnt/data/syncmark_pkg/proposal/COMP9993_Research_Proposal.docx')

ACCENT = RGBColor(26, 71, 122)
DARK = RGBColor(32, 32, 32)
MUTED = RGBColor(88, 96, 105)
LIGHT_BG = 'EDF3F8'
BORDER = 'D9E2EC'


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:%s' % edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn('w:%s' % key), str(edge_data[key]))


def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": BORDER, "space": 0},
                bottom={"val": "single", "sz": 6, "color": BORDER, "space": 0},
                left={"val": "single", "sz": 6, "color": BORDER, "space": 0},
                right={"val": "single", "sz": 6, "color": BORDER, "space": 0},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name='Aptos', size=11, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_styled_paragraph(doc, text, style=None, size=11, bold=False, color=DARK, space_after=6, left_indent=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
    normal.font.size = Pt(10.5)

    for name, size in [('Title', 22), ('Heading 1', 15), ('Heading 2', 12.5), ('Heading 3', 11.5)]:
        style = styles[name]
        style.font.name = 'Aptos'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.font.bold = True


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run('COMP9993 Research Proposal')
    set_font(r, size=10.5, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('SyncMark: Synchronization-Aware Multi-bit Watermarking')
    set_font(r, size=23, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run('for Short and Edited LLM Texts')
    set_font(r, size=18, bold=True, color=ACCENT)

    # separator line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('')
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'A6B5C5')
    pbdr.append(bottom)
    p_pr.append(pbdr)

    info = [
        ('Student', 'Zhan Ma (Master of Information Technology, UNSW Sydney)'),
        ('Project continuity', 'Built directly from the completed COMP9991 report and its negative findings'),
        ('Proposed 9993 direction', 'Autoregressive LLM text watermarking with synchronization-aware outer coding'),
        ('Primary ambition', 'A thesis-ready system and a paper-quality benchmark on short-text robustness'),
        ('Date', '13 April 2026'),
    ]

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.8), Inches(4.8)]
    for k, v in info:
        row = table.add_row()
        row.height = None
        for idx, txt in enumerate((k, v)):
            cell = row.cells[idx]
            cell.width = widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if idx == 0:
                rr = p.add_run(txt)
                set_font(rr, size=10.5, bold=True, color=ACCENT)
                set_cell_shading(cell, 'F7FAFC')
            else:
                rr = p.add_run(txt)
                set_font(rr, size=10.5)
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
    set_table_borders(table)

    doc.add_paragraph('')

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    cell = box.cell(0, 0)
    cell.width = Inches(6.6)
    set_cell_shading(cell, LIGHT_BG)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 10, "color": 'C9D6E2', "space": 0},
        bottom={"val": "single", "sz": 10, "color": 'C9D6E2', "space": 0},
        left={"val": "single", "sz": 10, "color": 'C9D6E2', "space": 0},
        right={"val": "single", "sz": 10, "color": 'C9D6E2', "space": 0},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('Executive summary')
    set_font(r, size=11.5, bold=True, color=ACCENT)
    summary = (
        'The proposal argues that the main failure mode of short-text multi-bit watermarking is not only raw bit corruption but also synchronization loss caused by insertion, deletion, and tokenization drift. '
        'It proposes SyncMark, an outer framing and alignment-decoding layer designed to sit on top of an existing inner watermark such as BiMark. '
        'The 9993 goal is to convert the negative findings of COMP9991 into a publishable positive contribution: a realistic benchmark and a synchronization-aware recovery method for short and edited LLM outputs.'
    )
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(summary)
    set_font(r, size=10.5)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_main_table(doc):
    add_heading(doc, 'Method overview', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ['Component', 'Role', 'Why it is needed', 'Connection to prior work']
    widths = [Inches(1.2), Inches(1.7), Inches(2.0), Inches(1.6)]
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, 'EAF1F7')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=10.2, bold=True, color=ACCENT)
    rows = [
        ('Inner watermark', 'Produce local bit evidence at each token', 'Keeps the system grounded in existing watermarking practice', 'BiMark is the preferred inner baseline'),
        ('Cycle-unique anchors', 'Re-lock the decoder after edits', 'Turns global drift into local damage', 'Direct response to insertion/deletion and char-level attacks'),
        ('Alignment decoder', 'Map noisy observed bits back to payload slots', 'Handles synchronization explicitly rather than implicitly', 'Brings coding-theoretic structure beyond naive ECC'),
        ('Checksum / validity test', 'Reject implausible payloads', 'Prevents fragile near-miss decodes from being over-counted', 'Compatible with any inner scheme'),
    ]
    for row in rows:
        tr = table.add_row()
        for i, text in enumerate(row):
            cell = tr.cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, size=9.9)
    set_table_borders(table)
    doc.add_paragraph('')


def add_experiment_table(doc):
    add_heading(doc, 'Planned experiment matrix', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ['Dimension', 'Primary settings', 'Main output', 'Decision criterion']
    widths = [Inches(1.2), Inches(2.2), Inches(1.5), Inches(1.7)]
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, 'EAF1F7')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=10.0, bold=True, color=ACCENT)
    rows = [
        ('Payload', '8 / 16 / 32 bits', 'Bit accuracy, exact recovery', 'Method must remain usable at 16-32 bits'),
        ('Length', '100 / 150 / 200 / 300 tokens', 'Short-text recovery curve', 'Main gains should appear in 100-300 tokens'),
        ('Attacks', 'token del/ins/sub, paraphrase, char-level', 'Robustness table', 'Character-level attacks must be in main results'),
        ('Quality', 'PPL and optional judge score', 'Quality-vs-robustness trade-off', 'Improvement must not come from severe quality collapse'),
        ('Baselines', 'BiMark, BiMark+ECC, XMark, DERMARK, MajorMark, MirrorMark', 'Comparative benchmark', 'At minimum beat BiMark+naive ECC convincingly'),
    ]
    for row in rows:
        tr = table.add_row()
        for i, text in enumerate(row):
            cell = tr.cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, size=9.8)
    set_table_borders(table)
    doc.add_paragraph('')


def add_timeline_table(doc):
    add_heading(doc, '12-week work plan', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.0), Inches(3.1), Inches(1.5)]
    headers = ['Weeks', 'Main task', 'Deliverable']
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, 'EAF1F7')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=10.0, bold=True, color=ACCENT)
    rows = [
        ('1-2', 'Clean BiMark reproduction, confirm 9991 baseline, run synthetic SyncMark validation', 'reproducible baseline + preview plots'),
        ('3-4', 'Implement outer framing and alignment decoder on top of a simple reference inner watermark', 'working end-to-end prototype'),
        ('5-7', 'Integrate SyncMark with BiMark inner layer; run token-level edit benchmarks', 'main robustness tables v1'),
        ('8-9', 'Run character-level and paraphrase attacks; perform ablations', 'attack suite + ablation section'),
        ('10-11', 'Compare against XMark/DERMARK/MajorMark/MirrorMark as compute permits', 'comparison plots and discussion'),
        ('12', 'Write thesis chapter / paper draft and prepare supervisor-facing summary', 'submission-ready draft'),
    ]
    for row in rows:
        tr = table.add_row()
        for i, text in enumerate(row):
            cell = tr.cells[i]
            cell.width = widths[i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i != 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, size=9.7)
    set_table_borders(table)
    doc.add_paragraph('')


def add_references(doc):
    add_heading(doc, 'Selected references', level=1)
    refs = [
        'Kirchenbauer et al. A Watermark for Large Language Models. ICML 2023.',
        'Hu et al. Unbiased Watermark for Large Language Models. 2023.',
        'Wu et al. A Resilient and Accessible Distribution-Preserving Watermark for Large Language Models. ICML 2024.',
        'Qu et al. Provably Robust Multi-bit Watermarking for AI-generated Text. USENIX Security 2025.',
        'Chao et al. Watermarking Language Models with Error Correcting Codes. WMARK@ICLR 2025 / TMLR under review in 2026.',
        'Feng et al. BiMark: Unbiased Multilayer Watermarking for Large Language Models. ICML 2025.',
        'Lin et al. DERMARK: A Dynamic, Efficient and Robust Multi-bit Watermark for Large Language Models. 2025.',
        'Xu et al. Majority Bit-Aware Watermarking for Large Language Models. 2025.',
        'Jiang et al. MirrorMark: A Distortion-Free Multi-Bit Watermark for Large Language Models. 2026.',
        'Xu et al. XMark: Reliable Multi-Bit Watermarking for LLM-Generated Texts. ACL 2026.',
        'Zhang et al. Character-Level Perturbations Disrupt LLM Watermarks. NDSS 2026.',
        'Huang et al. RLCracker: Exposing the Vulnerability of LLM Watermarks with Adaptive RL Attacks. 2025.',
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.03
        r = p.add_run(ref)
        set_font(r, size=9.5, color=MUTED)


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)

    add_heading(doc, '1. Motivation and problem statement', level=1)
    add_styled_paragraph(
        doc,
        'The completed COMP9991 stage established an important negative result: in short texts, simply layering ordinary ECC on top of a multi-bit watermark can reduce recovery rather than improve it. That observation should not be treated as failure. Instead, it should be treated as evidence that the dominant bottleneck was mis-specified.',
        size=10.7,
    )
    add_styled_paragraph(
        doc,
        'This proposal argues that post-generation edits should be modeled as a synchronization problem. Real edits do not only flip local bits; they also insert, delete, and retokenize parts of the text, causing downstream position drift. The 9993 project will therefore focus on a synchronization-aware outer layer for short-text multi-bit watermarking.',
        size=10.7,
    )

    add_heading(doc, '2. Core hypothesis', level=1)
    add_bullets(doc, [
        'Under insertion/deletion and character-level perturbations, exact message recovery fails largely because the decoder loses alignment between observed local signals and original bit positions.',
        'Redundancy is therefore more useful when some of it is allocated to synchronization anchors rather than only to ordinary parity bits.',
        'A synchronization-aware outer decoder should improve exact recovery in the 100-300 token regime while remaining compatible with existing inner watermarking schemes such as BiMark.',
    ])

    add_main_table(doc)

    add_heading(doc, '3. Proposed method: SyncMark', level=1)
    add_numbered(doc, [
        'Append a lightweight checksum to the payload message and repeat the protected payload across multiple short cycles.',
        'Insert a cycle-unique anchor sequence at the beginning of each cycle using a keyed pseudo-random generator.',
        'Use an inner watermark to bias token generation toward the next target bit in the anchor+payload stream.',
        'At detection time, first align the observed noisy bit stream to the expected layout using a dynamic-programming decoder with strong anchor rewards and gap penalties.',
        'Aggregate aligned payload observations by message index and recover the final message with majority voting plus checksum validation.',
    ])
    add_styled_paragraph(
        doc,
        'The key novelty is not a larger codeword. The novelty is that the outer code explicitly separates “recover alignment” from “recover payload”. This turns a single local edit from a global decoding disaster into a localized disturbance that the decoder can often contain.',
        size=10.7,
    )

    add_heading(doc, '4. Why this is better matched to the current literature', level=1)
    add_bullets(doc, [
        'BiMark provides continuity with the 9991 codebase and remains a strong short-text multi-bit baseline.',
        'DERMARK, MirrorMark, and XMark all signal that the field has moved toward finite-length capacity allocation, higher reliability, and stronger multi-bit design rather than naive payload expansion.',
        'Recent attack work shows that character-level perturbations and adaptive removal attacks must be treated as first-class robustness tests.',
        'Diffusion-language-model watermarking is advancing quickly, so it should be treated as a stretch goal rather than the main 9993 path unless the autoregressive study is already complete.',
    ])

    add_experiment_table(doc)

    add_heading(doc, '5. Risk assessment and mitigation', level=1)
    add_bullets(doc, [
        'Risk: gains appear only in synthetic experiments. Mitigation: keep the synthetic simulator as a design-checking tool, then move quickly to full LLM experiments.',
        'Risk: the reference inner watermark is too simple. Mitigation: treat the simple inner layer only as a scaffold and integrate the outer layer with BiMark early.',
        'Risk: quality degrades if every token is marked. Mitigation: the final thesis system should use a lower-distortion inner marker and compare under matched quality budgets.',
        'Risk: baselines are expensive to reproduce. Mitigation: prioritize BiMark, BiMark+naive ECC, and XMark first; add DERMARK/MajorMark/MirrorMark as compute permits.',
    ])

    add_heading(doc, '6. Expected contribution and publication path', level=1)
    add_styled_paragraph(
        doc,
        'A successful outcome would deliver three things: (i) a realistic short-text benchmark with insertion/deletion and character-level attacks, (ii) a synchronization-aware recovery mechanism that clearly beats naive ECC in the same regime, and (iii) a paper-ready empirical story explaining why short-text failure is driven by synchronization loss. The most realistic target is an ARR / Findings-style NLP submission, with watermarking or trustworthy-AI workshops as a backup path.',
        size=10.7,
    )

    add_timeline_table(doc)
    add_references(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f'[OK] wrote {OUT}')


if __name__ == '__main__':
    main()
